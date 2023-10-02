from docx.api import Document
import sys
import re
import psycopg2
from datetime import datetime

# connect to the database
conn = psycopg2.connect(
    database="StudentRestaurantUOI",
    host="localhost",
    user="postgres",
    password="admin",
    port="5432",
)


# meals dictionary
meals = {}

vowel = {
    "ά": "α",
    "έ": "ε",
    "ή": "η",
    "ί": "ι",
    "ό": "ο",
    "ύ": "υ",
    "ώ": "ω",
}

# Date data
year = 2023
month = 0
dayNumber = 0

# document to be read
document = None

# get the document name from the arguments
if len(sys.argv) > 2:
    if (sys.argv[1].endswith(".docx")):
        try:
            # open the document
            document = Document(sys.argv[1])
        except:
            # if the document is not found
            print("Error opening file")
            exit()
    else:
        # if the document is not a .docx file
        print("File should be a .docx file")
        exit()
    if (sys.argv[2].isdigit()):
        month = int(sys.argv[2])
        if month < 1 or month > 12:
            print("Month number should be between 1 and 12")
            exit()
    else:
        print("Month number should be a number")
        exit()
else:
    # if no document name is provided
    print("Please provide a document name and month number in the arguments")
    exit()


# read the meals from the database and parse them to meals dictionary
def read_meals():
    cursor = conn.cursor()
    postgreSQL_select_Query = "select * from meals"
    cursor.execute(postgreSQL_select_Query)
    meal_records = cursor.fetchall()
    for row in meal_records:
        meals[row[1]] = row[0]


# change the vowels without the tonous so that it's always the same
def compareString(string1, string2):
    string1 = string1.lower()
    string2 = string2.lower()
    string1 = string1.replace(" ", "")
    string2 = string2.replace(" ", "")
    for key in vowel:
        string1 = string1.replace(key, vowel[key])
        string2 = string2.replace(key, vowel[key])
    return string1 == string2


def printMyData(data):
    for day in data:
        for food in day:
            print(food)
        print("")


def existsInMeals(meal):
    for key in meals:
        if (compareString(meal, key)):
            return meals[key]
    return False


def sqlInsertMeal(meal):
    sql_insert = "INSERT INTO meals (name) VALUES ('" + meal+"');"
    cursor = conn.cursor()
    cursor.execute(sql_insert)
    conn.commit()
    cursor.close()
    meals[meal] = len(meals)+1
    return meals[meal]


def pushDayToDatabase(day):
    global meals
    date = day[0]
    date = datetime.strptime(date, "%d/%m/%Y")
    meal_type = day[1]
    first_dish = day[2]
    main_dish = day[3]
    extra_dish = day[4]
    side1_dish = day[5]
    side2_dish = day[6]

    result = existsInMeals(first_dish)
    if (result != False):
        first_dish = result
    else:
        first_dish = sqlInsertMeal(first_dish)

    result = existsInMeals(main_dish)
    if (result != False):
        main_dish = result
    else:
        main_dish = sqlInsertMeal(main_dish)

    result = existsInMeals(extra_dish)
    if (result != False):
        extra_dish = result
    else:
        extra_dish = sqlInsertMeal(extra_dish)

    result = existsInMeals(side1_dish)
    if (result != False):
        side1_dish = result
    else:
        side1_dish = sqlInsertMeal(side1_dish)

    result = existsInMeals(side2_dish)
    if (result != False):
        side2_dish = result
    else:
        side2_dish = sqlInsertMeal(side2_dish)

    sql = "INSERT INTO program (date, meal_type, first_dish, main_dish, special_dish, side_dish1, side_dish2) VALUES('"+str(date)+"', '" + \
        meal_type+"', '"+str(first_dish)+"', '"+str(main_dish)+"', '"+str(
        extra_dish)+"', '"+str(side1_dish)+"', '"+str(side2_dish)+"')"
    cursor = conn.cursor()
    cursor.execute(sql)
    conn.commit()
    cursor.close()


def pushToDatabase(data):
    show = input("Do you want to see the data? (y/n)")
    if (show == "y"):
        printMyData(data)
    verify = input("Do you want to push this data to the database? (y/n)")
    if (verify == "y"):
        read_meals()
        for day in data:
            pushDayToDatabase(day)
    else:
        exit()


def readTable(table):
    global dayNumber
    data = []
    # skip the first 2 columns that are two titles
    for column in list(table.columns)[2:]:
        # skip the first row that is the day of the week
        for i, cell in enumerate(column.cells[1:]):
            mealType = ''
            # if the cell is not empty skip
            if (cell.text != ""):
                if (i == 0):
                    dayNumber += 1
                    mealType = "meal"
                elif (i == 1):
                    mealType = "dinner"
                # so here cell.text is the real program for a day
                # remove double new lines and replace them with ~
                programString = cell.text.replace('\n\n', '~')
                programString = re.sub('\n+', ' ', programString)

                # remove double spaces
                programString = re.sub(' +', ' ', programString)

                # psudo strip
                programString = programString.replace('~ ', '~')
                programString = programString.replace(' ~', '~')

                # replace ~ή~ with ~
                programString = re.sub('~ή~', '~', programString)

                # so now there should be a clear programString with the food
                foodList = programString.split("~")

                # remove empty strings
                foodList = list(filter(None, foodList))

                # strip all the strings
                foodList = [x.strip() for x in foodList]

                if (len(foodList) != 5):
                    print("there is an Error with ", dayNumber)
                    continue

                # now we have a list with the food for the day
                # insert in start the date and the meal type
                global year
                global month
                foodList.insert(0,
                                str(dayNumber)+"/"+str(month)+"/"+str(year))
                foodList.insert(1, mealType)
                data.append(foodList)
    return data


def readDocument():
    tables = document.tables
    programData = []
    for table in tables:
        data = readTable(table)
        for day in data:
            programData.append(day)
    pushToDatabase(programData)


readDocument()
